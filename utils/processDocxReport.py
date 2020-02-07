# pip install python-docx datefinder
from docx import Document
import os
import datefinder
import re


org_report_folder = '/home/xiaosongw/Downloads/anonymized/'
output_report_folder = '/home/xiaosongw/Downloads/anonymized_output/'

if not os.path.exists(output_report_folder):
    os.mkdir(output_report_folder)

# list all the reports in the folder
org_report_list = [file for file in os.listdir(org_report_folder) if file.endswith('.docx')]

report_index = []
for r_i, report_name in enumerate(org_report_list):
    f = open(os.path.join(org_report_folder, report_name), 'rb')
    document = Document(f)
    f.close()
    para_first = document.paragraphs[0]
    if len(para_first.text.split()) < 2:
        para_first.text = 'XXXX'

    for p_i, para in enumerate(document.paragraphs):
        matches = list(datefinder.find_dates(para.text, source=True, index=True))
        if len(matches) > 0:
            for date_t in matches:
                match_date = re.search(r'\d{1,2}/\d{1,2}/\d{2,4}', date_t[1])
                if match_date:
                    document.paragraphs[p_i].text = document.paragraphs[p_i].text.replace(date_t[1], 'XX/XX/XXXX')
                    print(date_t[1]+' in '+ report_name)
        if 'Dictated By:' in para.text:
            para.text = 'Dictated by: XXXX'

    # print(document)
    document.save(os.path.join(output_report_folder, '%05d'%r_i +'.docx'))

    report_index.append('%05d'%r_i + '.docx;' + report_name + '\n')

fh = open(os.path.join(output_report_folder, 'report_index.csv'),'w')
fh.writelines(report_index)
fh.close()